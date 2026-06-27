#!/usr/bin/env python3
"""Generate a ~10-page Word document of CURATED KEY EXCERPTS from the 7-DOF
teach-pendant source (pendant app + motion backend).

Each excerpt is verbatim within its selected line ranges; code omitted between
ranges is marked with ``# ...``. Plain monospace, no line numbers, light-grey
code blocks. The per-file caption records the exact source line ranges so the
provenance of every excerpt is traceable.

(The complete, full-file version of every file lives in
``make_code_appendix_full.py``.)
"""
from __future__ import annotations

import os
from docx import Document
from docx.shared import Pt, Cm, Mm, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

WS = "/home/fish/7dof_ws"
OUT = os.path.join(WS, "Teach_Pendant_Source_Code.docx")

CODE_FILL = "F2F2F2"
CODE_FONT = "Consolas"
CODE_SIZE = 8.5

THESIS_TITLE = "Development of a Software Teach Pendant for a 7-DOF Robot Arm"

# ---------------------------------------------------------------------------
# Excerpt manifest. Each file lists 1-based inclusive line ranges to include;
# gaps between ranges are rendered as a "# ..." elision marker.
# ---------------------------------------------------------------------------
SECTION1 = (
    "1  Teach Pendant Application",
    "The teach pendant is a standalone Qt (PyQt6) desktop application that "
    "couples to the robot only through a small set of ROS 2 topics. The "
    "excerpts below show the signature parts of each module; routine setup and "
    "boilerplate are elided with “# ...”.",
    [
        ("teach_pendant/pendant7dof/ros_bridge.py",
         "ROS Bridge — Interface Contract and Cartesian Jog",
         "The single rclpy node owned by the GUI: the joint-name/limit contract, the publishers and subscribers for the contract topics, and the workspace-clamped Cartesian jog that keeps the end-effector target inside the reachable dome.",
         [(29, 53), (62, 72), (121, 132), (156, 166), (381, 404), (427, 435), (460, 461)]),
        ("teach_pendant/pendant7dof/gui/main_window.py",
         "Main Window — Jog Dispatch and Simulation Control",
         "How the GUI dispatches user input: the joint/Cartesian mode and joint-group toggles, the joystick callback that streams jog commands to the bridge, and the Simulation on/off control that launches the ROS backend.",
         [(1076, 1095), (1799, 1849)]),
        ("teach_pendant/pendant7dof/gui/joystick.py",
         "Jog Joystick Widget",
         "The custom two-axis jog knob: per-drag axis-lock (slider feel for Cartesian jog), the fixed-rate velocity-style stream to the consumer, and the final zero tick on release (paint methods omitted).",
         [(28, 61), (258, 298)]),
        ("teach_pendant/pendant7dof/gui/drawing_canvas.py",
         "Drawing Canvas Widget",
         "Stroke capture on the millimetre-scaled canvas (strokes hold rather than trail the edge when the mouse leaves), the JSON drawing handed to the planner, and the live pen-tip dot that fades when the pen lifts.",
         [(74, 92), (123, 174)]),
    ],
)

SECTION2 = (
    "2  Motion Backend Nodes",
    "The pendant launches and drives a set of ROS 2 nodes that perform the "
    "kinematics and trajectory generation, all built directly from the latched "
    "/robot_description rather than an idealised DH model. They communicate "
    "with the pendant only through the contract topics.",
    [
        ("src/arm_bot/arm_bot/ik_lib.py",
         "Kinematics Library — URDF Jacobian and DLS Solvers",
         "Shared kinematics: forward kinematics and the 6xN geometric Jacobian built from the URDF, and the interface and documented design of the task-priority pen-tip solver (position primary, pen-perpendicular orientation solved in the null space) used for tilt-capable drawing.",
         [(62, 64), (116, 143), (290, 319)]),
        ("src/arm_bot/arm_bot/ik_arm_v3.py",
         "Inverse-Kinematics Node — Real-Time Control Loop",
         "The 200 Hz damped-least-squares control tick: adaptive damping, a singularity-robust damping floor from the manipulability measure, weighted DLS, and self-fading null-space joint-centering and velocity damping.",
         [(240, 298), (306, 334)]),
        ("src/arm_bot/arm_bot/fk_arm_v3.py",
         "Forward-Kinematics Node",
         "Loads the latched /robot_description (TRANSIENT_LOCAL QoS), builds the chain, and republishes the end-effector pose on /ee_pose for every /joint_states update.",
         [(142, 165), (180, 210)]),
        ("src/arm_bot/arm_bot/drawing_batch_planner.py",
         "Drawing Batch Planner — Corner-Aware Waypoints",
         "Corner detection and per-segment resampling that keep genuine corners sharp (a single spline would round them), the corner-aware stroke splitting in the waypoint builder, and the pen-down draw loop.",
         [(1010, 1060), (1099, 1120), (1184, 1199), (1239, 1241)]),
        ("src/arm_bot/launch/pendant_backend.launch.py",
         "Backend Launch Graph",
         "The nodes the “Simulation ON” graph spawns and wires together — the IK node (with its soft joint limits), the FK node, and the IK-to-trajectory bridge (the drawing planner is spawned alongside).",
         [(139, 153), (162, 175)]),
    ],
)

SECTIONS = [SECTION1, SECTION2]


# ---------------------------------------------------------------------------
# helpers
# ---------------------------------------------------------------------------
def set_style_fonts(style, name):
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), name)


def shade(paragraph, fill=CODE_FILL):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def _emit_code_line(doc, text):
    p = doc.add_paragraph(style="CodeBlock")
    shade(p)
    run = p.add_run(text)
    t = run._r.find(qn("w:t"))
    if t is not None:
        t.set(qn("xml:space"), "preserve")


def _indent_of(line: str) -> str:
    return line[: len(line) - len(line.lstrip(" "))]


def add_excerpt(doc, lines, ranges):
    """Emit the given 1-based inclusive line ranges, with a '# ...' elision
    marker between non-adjacent ranges (indented to match the following code)."""
    ranges = sorted(ranges)
    for ri, (a, b) in enumerate(ranges):
        if ri > 0:
            prev_b = ranges[ri - 1][1]
            if a > prev_b + 1:                       # something was skipped
                first = lines[a - 1] if a - 1 < len(lines) else ""
                _emit_code_line(doc, _indent_of(first) + "# ...")
        for ln in range(a, b + 1):
            _emit_code_line(doc, lines[ln - 1].expandtabs(4))


def fmt_ranges(ranges):
    return ", ".join(f"{a}–{b}" if a != b else f"{a}" for a, b in sorted(ranges))


def add_page_number_footer(section):
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    f2 = OxmlElement("w:fldChar"); f2.set(qn("w:fldCharType"), "end")
    run._r.append(f1); run._r.append(instr); run._r.append(f2)
    run.font.size = Pt(9)
    run.font.name = "Calibri"


# ---------------------------------------------------------------------------
# build
# ---------------------------------------------------------------------------
doc = Document()
doc.core_properties.title = "Source Code — Key Excerpts"
doc.core_properties.subject = THESIS_TITLE
doc.core_properties.author = "Chhay-Yeah"

sec = doc.sections[0]
sec.page_width = Mm(210)
sec.page_height = Mm(297)
sec.top_margin = Cm(1.5)
sec.bottom_margin = Cm(1.5)
sec.left_margin = Cm(1.6)
sec.right_margin = Cm(1.6)
add_page_number_footer(sec)

for hname in ("Heading 1", "Heading 2", "Title"):
    try:
        doc.styles[hname].font.color.rgb = RGBColor(0, 0, 0)
    except KeyError:
        pass

code_style = doc.styles.add_style("CodeBlock", WD_STYLE_TYPE.PARAGRAPH)
code_style.base_style = doc.styles["Normal"]
code_style.font.name = CODE_FONT
code_style.font.size = Pt(CODE_SIZE)
code_style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
set_style_fonts(code_style, CODE_FONT)
pf = code_style.paragraph_format
pf.space_before = Pt(0)
pf.space_after = Pt(0)
pf.line_spacing = 1.0
pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
pf.left_indent = Cm(0.25)
pf.widow_control = False

# Compact front matter (no separate title page / TOC — keeps it to ~10 pages).
title = doc.add_heading("Source Code — Key Excerpts", level=0)
title.paragraph_format.space_after = Pt(2)
sub = doc.add_paragraph()
r = sub.add_run("Appendix — " + THESIS_TITLE)
r.italic = True
r.font.size = Pt(11)
intro = doc.add_paragraph()
intro.add_run(
    "This appendix collects the key source-code excerpts of the teach pendant, "
    "grouped into the pendant application and the ROS 2 motion-backend nodes it "
    "drives. Each excerpt is reproduced verbatim from the repository; the path "
    "and the exact source line numbers are given beneath every heading, and "
    "omitted code is shown as “# ...”."
)
intro.paragraph_format.space_after = Pt(8)

total = 0
for sec_title, sec_intro, files in SECTIONS:
    h1 = doc.add_heading(sec_title, level=1)
    h1.paragraph_format.space_before = Pt(8)
    h1.paragraph_format.space_after = Pt(4)
    si = doc.add_paragraph()
    si.add_run(sec_intro)
    si.paragraph_format.space_after = Pt(8)

    for relpath, heading, desc, ranges in files:
        lines = open(os.path.join(WS, relpath), encoding="utf-8").read().split("\n")
        nsel = sum(b - a + 1 for a, b in ranges)
        total += nsel

        h2 = doc.add_heading(heading, level=2)
        h2.paragraph_format.keep_with_next = True
        h2.paragraph_format.space_before = Pt(8)
        h2.paragraph_format.space_after = Pt(2)

        cap = doc.add_paragraph()
        cap.paragraph_format.keep_with_next = True
        cap.paragraph_format.space_after = Pt(4)
        cpath = cap.add_run(relpath)
        cpath.font.name = "Consolas"; cpath.font.size = Pt(9); cpath.bold = True
        cdesc = cap.add_run("   ·   " + desc)
        cdesc.italic = True; cdesc.font.size = Pt(9)
        cprov = cap.add_run("   [excerpt — source lines " + fmt_ranges(ranges) + "]")
        cprov.font.name = "Consolas"; cprov.font.size = Pt(8)
        cprov.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        add_excerpt(doc, lines, ranges)

doc.save(OUT)
print("WROTE", OUT)
print("excerpt code lines:", total)
