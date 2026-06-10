#!/usr/bin/env python3
"""
make_table_4_2.py  —  builds Table 4.2 (objectives vs outcomes) as a .docx.

This table is AUTHORED, not measured — no rosbag needed. Edit the ROWS list
below: replace each objective with your exact Section 1.3 wording. The outcome
/ evidence cells are already filled with the measured results from this run
(Figs 4.1–4.5, Table 4.1).

    pip install --user python-docx        # one-time
    python3 src/arm_bot/analysis/make_table_4_2.py --out table_4_2.docx
"""
import argparse

# Each row: (objective, outcome that demonstrates it, evidence)
# Objectives are the verbatim Section 1.3 objectives of the thesis (paras 93-96);
# outcomes are the measured results from the Chapter 4 drawing run.
ROWS = [
    ("Objective 1 — Implement the core pendant feature set in software: joint and "
     "Cartesian jogging, waypoint teaching and editing, motion-sequence authoring and "
     "execution, and live state monitoring, as a self-contained PyQt6 desktop application.",
     "All five operator modes are implemented in a self-contained PyQt6 application: "
     "joint/Cartesian jogging, waypoint teaching and editing, flowchart-based "
     "motion-sequence authoring and execution, live state monitoring, and configuration.",
     "Chapter 3 (interface and modes); published application"),
    ("Objective 2 — Integrate the pendant with ROS 2 so operator actions produce "
     "coherent motion: a ROS 2 client that subscribes to robot state, invokes the lab's "
     "FK/IK nodes, and dispatches motion through the JointTrajectoryController.",
     "The in-process ROS 2 client drives the arm coherently — subscribing to "
     "/joint_states, invoking the laboratory's FK/IK nodes, and dispatching motion "
     "through the JointTrajectoryController; commanded operator paths produce the "
     "intended motion on the simulated arm.",
     "Figures 4.1, 4.3, 4.4; ROS 2 bridge (Chapter 3)"),
    ("Objective 3 — Implement a Cartesian drawing mode as an end-to-end demonstration "
     "exercising the full pipeline and giving a visually verifiable benchmark of correctness.",
     "The drawing mode exercises the complete pipeline (operator input → offline batch "
     "path planning → dispatch → execution); the executed end-effector path follows the "
     "commanded path to RMS 0.11 mm (max 0.61 mm), with smooth joint motion and joint 6 "
     "remaining within its limits throughout.",
     "Figures 4.1, 4.2, 4.3, 4.4, 4.5; Table 4.1"),
    ("Objective 4 — Validate and package the pendant in Gazebo Ignition simulation and "
     "distribute it via PyPI for Linux with ROS 2 Humble, with real-hardware changes "
     "identified in the conclusion.",
     "The pendant is validated end-to-end in Gazebo Ignition (via ign_ros2_control) and "
     "packaged as a pip-installable application published to PyPI; the changes required "
     "for real-hardware operation are identified in Chapter 5.",
     "Figure 4.2; PyPI release; Chapter 5 (future work)"),
]


def build(out_path):
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    doc.add_heading('Table 4.2 — Objectives vs Outcomes', level=2)

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    for c, text in zip(hdr, ('Objective (Section 1.3)', 'Outcome', 'Evidence')):
        c.text = ''
        run = c.paragraphs[0].add_run(text)
        run.bold = True

    for obj, outcome, evidence in ROWS:
        cells = table.add_row().cells
        cells[0].text = obj
        cells[1].text = outcome
        cells[2].text = evidence

    # modest fixed column widths
    from docx.shared import Inches
    widths = (Inches(2.4), Inches(3.2), Inches(1.2))
    for row in table.rows:
        for cell, w in zip(row.cells, widths):
            cell.width = w

    doc.save(out_path)
    print(f'wrote {out_path}')


if __name__ == '__main__':
    ap = argparse.ArgumentParser()
    ap.add_argument('--out', default='table_4_2.docx')
    args = ap.parse_args()
    build(args.out)
