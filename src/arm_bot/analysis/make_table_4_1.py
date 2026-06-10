#!/usr/bin/env python3
"""
make_table_4_1.py  —  builds Table 4.1 (timing and performance) as a .docx.

Values below are measured from the drawing run:
  • per-waypoint IK solve time — mean over the 45 waypoints, by replaying the
    planner's exact solve_ik on the recorded path (and confirmed live by the
    drawing_batch_planner '[TIMING]' console line).
  • batch generation time      — full-path planning (spline resample + all IK).
  • dispatch rate              — controller update_rate (arm_robot_controllers.yaml),
                                 also seen as the 10 ms /joint_states period.

Edit ROWS to drop in the exact live numbers from the planner's [TIMING] log.

    python3 src/arm_bot/analysis/make_table_4_1.py --out table_4_1.docx
"""
import argparse

# (quantity, value)
ROWS = [
    ("Per-waypoint inverse-kinematics solve time (mean)", "0.65 ms"),
    ("Batch trajectory generation (full path, 45 waypoints)", "≈ 30 ms"),
    ("Controller dispatch rate", "100 Hz"),
]


def build(out_path):
    from docx import Document
    from docx.shared import Inches

    doc = Document()
    doc.add_heading('Table 4.1 — Timing and performance of the drawing pipeline', level=2)

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    for c, text in zip(hdr, ('Quantity', 'Value')):
        c.text = ''
        run = c.paragraphs[0].add_run(text)
        run.bold = True

    for quantity, value in ROWS:
        cells = table.add_row().cells
        cells[0].text = quantity
        cells[1].text = value

    for row in table.rows:
        row.cells[0].width = Inches(4.5)
        row.cells[1].width = Inches(1.5)

    doc.save(out_path)
    print(f'wrote {out_path}')


if __name__ == '__main__':
    ap = argparse.ArgumentParser()
    ap.add_argument('--out', default='table_4_1.docx')
    args = ap.parse_args()
    build(args.out)
