#!/usr/bin/env python3
"""Generate a Word document containing the relevant source code of the
7-DOF teach-pendant application, with a heading per file.

Scope: pendant app (teach_pendant/pendant7dof/) + motion backend nodes
(src/arm_bot/...). Plain monospace, no line numbers, light-grey code blocks.
"""
from __future__ import annotations

import os
from docx import Document
from docx.shared import Pt, Cm, Mm, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

WS = "/home/fish/7dof_ws"
OUT = os.path.join(WS, "Teach_Pendant_Source_Code.docx")

CODE_FILL = "F2F2F2"      # light grey behind code
CODE_FONT = "Consolas"
CODE_SIZE = 8.5           # pt

THESIS_TITLE = "Development of a Software Teach Pendant for a 7-DOF Robot Arm"

# ---------------------------------------------------------------------------
# File manifest: (relative path, heading title, one-line description)
# ---------------------------------------------------------------------------
SECTION1 = (
    "1  Teach Pendant Application",
    "The teach pendant is a standalone Qt (PyQt6) desktop application, "
    "distributed as the pip package 7dof-pendant. It imports no robot node: it "
    "couples to the workspace purely through a ROS 2 name/type contract and "
    "bootstraps a colcon build of the workspace on launch. The modules below "
    "implement the application — its packaging and entry points, the "
    "environment bootstrap, the Qt GUI widgets, the ROS bridge that realises the "
    "interface contract, and a small on-disk settings store.",
    [
        ("teach_pendant/pyproject.toml",
         "Packaging",
         "Project metadata, dependencies and console entry points for the pip-installable 7dof-pendant package."),
        ("teach_pendant/pendant7dof/__init__.py",
         "Package Marker",
         "Top-level package marker and version string."),
        ("teach_pendant/pendant7dof/cli.py",
         "Command-Line Interface",
         "Command-line front end (launch / build / bundle / release / doctor sub-commands) for the 7dof-pendant executable."),
        ("teach_pendant/pendant7dof/bootstrap.py",
         "Environment Bootstrap",
         "Locates ROS 2 and the colcon workspace and re-execs the process inside a fully sourced shell so that rclpy is importable."),
        ("teach_pendant/pendant7dof/launcher.py",
         "Backend Launcher",
         "Manages the ROS 2 backend (pendant_backend.launch.py) as a child process started and torn down together with the GUI."),
        ("teach_pendant/pendant7dof/gui/__init__.py",
         "GUI Sub-package Marker",
         "Marker for the gui sub-package."),
        ("teach_pendant/pendant7dof/gui/app.py",
         "GUI Entry Point",
         "Initialises rclpy and Qt, creates the bridge node, main window and backend process, and pumps rclpy from the Qt event loop."),
        ("teach_pendant/pendant7dof/gui/main_window.py",
         "Main Application Window",
         "The central GUI orchestrator: Joint / Cartesian / Drawing mode tabs, joint sliders, jog joystick, drawing canvas, saved targets and tasks, the status panel and the Simulation on/off control."),
        ("teach_pendant/pendant7dof/gui/joystick.py",
         "Jog Joystick Widget",
         "A custom two-axis jog knob with mode-dependent cardinal labels, the joint-group toggle, per-drag axis-lock and the twist ring with its read-only legend."),
        ("teach_pendant/pendant7dof/gui/drawing_canvas.py",
         "Drawing Canvas Widget",
         "Captures freehand strokes on the fixed 100x100 mm canvas and renders the live pen-tip tracking dot, which fades when the pen lifts."),
        ("teach_pendant/pendant7dof/ros_bridge.py",
         "ROS Bridge",
         "The single rclpy node embedded in the GUI: holds the joint-name/limit contract, publishes and subscribes the contract topics, and implements jogging, the workspace clamp, preset motion and the drawing dispatch."),
        ("teach_pendant/pendant7dof/store.py",
         "Settings Store",
         "On-disk JSON persistence for saved jogging targets and motion tasks, under the user's config directory."),
    ],
)

SECTION2 = (
    "2  Motion Backend Nodes",
    "The pendant launches and drives a set of ROS 2 nodes that perform the "
    "kinematics and trajectory generation. They are spawned by "
    "pendant_backend.launch.py (the “Simulation ON” graph) and "
    "communicate with the pendant only through the contract topics: the IK node "
    "consumes /ee_target and emits /joint_commands, the FK node publishes "
    "/ee_pose, the drawing planner consumes drawing/strokes and emits a batch "
    "trajectory, and all kinematics are built directly from the latched "
    "/robot_description rather than an idealised DH model.",
    [
        ("src/arm_bot/arm_bot/ik_lib.py",
         "Kinematics Library",
         "Parses the URDF chain, builds FK and the 6xN geometric Jacobian, and implements the damped-least-squares solvers solve_ik (full pose) and solve_ik_tip (position-primary, pen tip)."),
        ("src/arm_bot/arm_bot/fk_arm_v3.py",
         "Forward-Kinematics Node",
         "Publishes /ee_pose from /joint_states using the latched /robot_description chain; the end-effector pose source of truth."),
        ("src/arm_bot/arm_bot/ik_arm_v3.py",
         "Inverse-Kinematics Node",
         "Builds FK and the geometric Jacobian from /robot_description and drives /joint_commands at 200 Hz with adaptive-damping DLS, null-space joint-centering and a per-step magnitude cap."),
        ("src/arm_bot/arm_bot/ik_to_trajectory.py",
         "IK-to-Trajectory Bridge",
         "Converts the IK node's /joint_commands stream into JointTrajectory goals for arm_controller."),
        ("src/arm_bot/arm_bot/drawing_batch_planner.py",
         "Drawing Batch Planner",
         "Maps normalised strokes onto the paper plane and solves pen-tip IK per waypoint with corner-aware resampling, per-stroke and corner dwells, pen-offset Z-compensation, table-tilt compensation and joint-speed retiming, emitting one batch JointTrajectory."),
        ("src/arm_bot/launch/pendant_backend.launch.py",
         "Backend Launch Graph",
         "The “Simulation ON” launch graph: brings up Gazebo (or the MoveIt demo), RViz, the controller spawners and the IK / FK / ik_to_trajectory / drawing_batch_planner nodes with their tuned parameters."),
    ],
)

SECTIONS = [SECTION1, SECTION2]


# ---------------------------------------------------------------------------
# helpers
# ---------------------------------------------------------------------------
def set_style_fonts(style, name):
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), name)


def shade(paragraph, fill=CODE_FILL):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def add_code_block(doc, text):
    """One paragraph per source line so the listing page-breaks cleanly while
    the grey background stays continuous."""
    if text.endswith("\n"):           # drop the single phantom trailing newline
        text = text[:-1]
    for line in text.split("\n"):
        line = line.expandtabs(4)
        p = doc.add_paragraph(style="CodeBlock")
        shade(p)
        run = p.add_run(line)
        t = run._r.find(qn("w:t"))
        if t is not None:
            t.set(qn("xml:space"), "preserve")


def add_toc(doc):
    para = doc.add_paragraph()
    run = para.add_run()
    f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-2" \\h \\z \\u'
    f2 = OxmlElement("w:fldChar"); f2.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Right-click here and choose “Update Field” (or select all and press F9) to build the table of contents."
    f3 = OxmlElement("w:fldChar"); f3.set(qn("w:fldCharType"), "end")
    for el in (f1, instr, f2, placeholder, f3):
        run._r.append(el)


def add_page_number_footer(section):
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    f2 = OxmlElement("w:fldChar"); f2.set(qn("w:fldCharType"), "end")
    run._r.append(f1); run._r.append(instr); run._r.append(f2)
    run.font.size = Pt(9)
    run.font.name = "Calibri"


# ---------------------------------------------------------------------------
# build
# ---------------------------------------------------------------------------
doc = Document()

# Core properties
doc.core_properties.title = "Source Code Listings"
doc.core_properties.subject = THESIS_TITLE
doc.core_properties.author = "Chhay-Yeah"

# Page geometry (A4, narrow side margins so code fits)
sec = doc.sections[0]
sec.page_width = Mm(210)
sec.page_height = Mm(297)
sec.top_margin = Cm(2.0)
sec.bottom_margin = Cm(2.0)
sec.left_margin = Cm(1.6)
sec.right_margin = Cm(1.6)
add_page_number_footer(sec)

# Neutral (black) heading colours for a formal appendix look
for hname in ("Heading 1", "Heading 2", "Title"):
    try:
        doc.styles[hname].font.color.rgb = RGBColor(0, 0, 0)
    except KeyError:
        pass

# CodeBlock paragraph style
code_style = doc.styles.add_style("CodeBlock", WD_STYLE_TYPE.PARAGRAPH)
code_style.base_style = doc.styles["Normal"]
code_style.font.name = CODE_FONT
code_style.font.size = Pt(CODE_SIZE)
code_style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
set_style_fonts(code_style, CODE_FONT)
pf = code_style.paragraph_format
pf.space_before = Pt(0)
pf.space_after = Pt(0)
pf.line_spacing = 1.0
pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
pf.left_indent = Cm(0.25)
pf.right_indent = Cm(0.0)
pf.widow_control = False

# ---- Title block ----
title = doc.add_heading("Source Code Listings", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Appendix — " + THESIS_TITLE)
r.italic = True
r.font.size = Pt(12)

doc.add_paragraph()
intro = doc.add_paragraph()
intro.add_run(
    "This appendix reproduces, verbatim, the source code central to the teach "
    "pendant. Each listing carries a heading and, beneath it, the file's path "
    "in the repository together with a one-line summary of its role. The code "
    "is grouped into two parts: the pendant application itself, and the ROS 2 "
    "motion-backend nodes that the pendant launches and drives. The two are "
    "deliberately decoupled — the pendant communicates with the robot only "
    "through a small set of ROS topics, never by importing a robot node."
)

doc.add_paragraph()
toc_heading = doc.add_paragraph()
tr = toc_heading.add_run("Contents")
tr.bold = True
tr.font.size = Pt(13)
add_toc(doc)

# ---- Sections ----
total_lines = 0
for sec_title, sec_intro, files in SECTIONS:
    h1 = doc.add_heading(sec_title, level=1)
    h1.paragraph_format.page_break_before = True
    h1.paragraph_format.space_after = Pt(6)

    si = doc.add_paragraph()
    si.add_run(sec_intro).italic = False
    si.paragraph_format.space_after = Pt(10)

    for relpath, heading, desc in files:
        abspath = os.path.join(WS, relpath)
        with open(abspath, encoding="utf-8") as fh:
            text = fh.read()
        nlines = len(text.rstrip("\n").split("\n")) if text.strip() else 0
        total_lines += nlines

        h2 = doc.add_heading(heading, level=2)
        h2.paragraph_format.keep_with_next = True
        h2.paragraph_format.space_before = Pt(14)
        h2.paragraph_format.space_after = Pt(2)

        cap = doc.add_paragraph()
        cap.paragraph_format.keep_with_next = True
        cap.paragraph_format.space_after = Pt(4)
        cpath = cap.add_run(relpath)
        cpath.font.name = "Consolas"
        cpath.font.size = Pt(9)
        cpath.bold = True
        cdesc = cap.add_run("   ·   " + desc + f"   ({nlines} lines)")
        cdesc.italic = True
        cdesc.font.size = Pt(9)

        add_code_block(doc, text)

doc.save(OUT)
print("WROTE", OUT)
print("total code lines embedded:", total_lines)
npar = len(doc.paragraphs)
print("paragraphs:", npar)
